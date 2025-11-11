# src/core/orchestrator.py

import os
from datetime import datetime
from dotenv import load_dotenv
from docx import Document # <-- Nova importa√ß√£o

from src.services import openai_service
from src.services.openai_service import transcrever_audio
from src.services.database_operations import get_text_from_file
from src.config import TEMPLATES_PATH

load_dotenv()

# --- FUN√á√ÉO PRINCIPAL DE AN√ÅLISE (sem altera√ß√µes) ---
def run_analysis_and_generate_artifacts(transcription: str, info_cliente: dict, doc_paths: list, progress_callback=None) -> tuple[str, dict]:
    # (Esta fun√ß√£o permanece a mesma)
    def report_progress(message):
        if progress_callback: progress_callback(message)
    context_docs = ""
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        error_msg = "ERRO: Chave da API da OpenAI n√£o encontrada."
        report_progress(f"-> [‚ùå] {error_msg}")
        return "", {}
    if doc_paths:
        report_progress(f"-> [‚öôÔ∏è] Processando {len(doc_paths)} documento(s) de apoio...")
        AUDIO_EXTENSIONS = ['.ogg', '.mp3', '.m4a', '.wav', '.mp4', '.mpeg']
        for path in doc_paths:
            try:
                _, file_extension = os.path.splitext(path)
                if file_extension.lower() in AUDIO_EXTENSIONS:
                    context_docs += f"--- TRANSCRI√á√ÉO DO √ÅUDIO: {os.path.basename(path)} ---\n"
                    context_docs += transcrever_audio(openai_api_key, path)
                else:
                    context_docs += f"--- CONTE√öDO DO DOCUMENTO: {os.path.basename(path)} ---\n"
                    context_docs += get_text_from_file(path)
                context_docs += "\n\n"
            except Exception as e:
                report_progress(f"-> [‚ö†Ô∏è] Falha ao processar o arquivo {path}: {e}")
    report_progress("-> [ü§ñ] IA est√° analisando e modelando o processo BPMN. Isso pode levar um momento...")
    analysis_result = openai_service.generate_business_analysis(
        api_key=openai_api_key, full_transcription=transcription,
        client_info=info_cliente, context_docs=context_docs
    )
    report_progress("-> [‚úÖ] An√°lise e modelagem da IA conclu√≠das.")
    bpmn_xml = analysis_result.get("bpmn_xml", "")
    spec_content = analysis_result.get("specification_content", {})
    return bpmn_xml, spec_content


# --- GERADOR DE DOCUMENTO .TXT (ATUALIZADO) ---
def generate_specification_document(spec_content: dict, client_name: str) -> str:
    try:
        template_path = os.path.join(TEMPLATES_PATH, "functional_specification_template.txt")
        if not os.path.exists(template_path):
            return "ERRO: O arquivo 'functional_specification_template.txt' n√£o foi encontrado."
        with open(template_path, "r", encoding="utf-8") as f:
            template = f.read()
        spec_content['date'] = datetime.now().strftime("%d/%m/%Y")
        for key, placeholder in [
            ("system_name", "{system_name}"), ("document_name", "{document_name}"),
            ("importance", "{importance}"), ("project_code", "{project_code}"),
            ("date", "{date}"), ("document_objective", "{document_objective}"),
            ("user_flow", "{user_flow}"), ("user_profiles", "{user_profiles}"),
            ("prototype_link", "{prototype_link}")
        ]:
            value = spec_content.get(key, placeholder)
            template = template.replace(placeholder, str(value))
        user_stories_value = spec_content.get("user_stories")
        if isinstance(user_stories_value, list):
            formatted_stories = "\n".join(
                f"- {story.get('description', 'N/A')}" for story in user_stories_value
            )
            template = template.replace("{user_stories}", formatted_stories)
        elif isinstance(user_stories_value, str):
            template = template.replace("{user_stories}", user_stories_value)
        else:
            template = template.replace("{user_stories}", "{user_stories}")
        functionalities_text = ""
        functionalities_list = spec_content.get('functionalities', [])
        if functionalities_list:
            for i, func in enumerate(functionalities_list):
                functionalities_text += f"6.{i+1} {func.get('title', 'Funcionalidade sem t√≠tulo')}\n"
                functionalities_text += f"----------------------------------------\n"
                functionalities_text += f"Descri√ß√£o: {func.get('description', 'N/A')}\n"
                functionalities_text += f"Acionador: {func.get('trigger', 'N/A')}\n"
                functionalities_text += f"Integra√ß√µes: {func.get('integrations', 'N/A')}\n"
                functionalities_text += f"V√≠nculo com Telas: {func.get('screen_links', 'N/A')}\n"
                functionalities_text += f"Campos: {func.get('fields', 'N/A')}\n\n"
                functionalities_text += f"Requisitos Funcionais:\n"
                rf_list = func.get('functional_requirements', [])
                if rf_list:
                    # Formata a lista de requisitos funcionais
                    for req in rf_list:
                        functionalities_text += f"- {req}\n"
                else:
                    functionalities_text += "- Nenhum requisito funcional detalhado.\n"
                functionalities_text += "\n\n"
        template = template.replace("{functionalities_section}", functionalities_text or "Nenhuma funcionalidade detalhada foi gerada.")
        return template
    except Exception as e:
        return f"Erro ao gerar o documento de especifica√ß√£o: {e}"


# --- NOVA FUN√á√ÉO: GERADOR DE DOCUMENTO .DOCX (VERS√ÉO CORRIGIDA) ---
def generate_word_document(spec_content: dict):
    """
    Preenche um template do Word (.docx) com o conte√∫do gerado pela IA.
    """
    try:
        template_path = os.path.join(TEMPLATES_PATH, "EF-Template-(MOSTEN).docx")
        if not os.path.exists(template_path):
            raise FileNotFoundError("O template 'EF-Template-(MOSTEN).docx' n√£o foi encontrado na pasta 'templates'.")
        
        doc = Document(template_path)
        
        # Dicion√°rio de placeholders para substitui√ß√£o
        replacements = {
            '[DESCRI√á√ÉO DO OBJETIVO DO DOCUMENTO]': str(spec_content.get("document_objective", "")),
            '[DESCRI√á√ÉO DA HIST√ìRIA DO USU√ÅRIO]': str(spec_content.get("user_stories", "")),
            '[FLUXOGRAMA COM A JORNADA DO USU√ÅRIO]': str(spec_content.get("user_flow", "")),
            '[INDICAR PERFIS DE USU√ÅRIOS OU ATORES QUE UTILIZAR√ÉO O SISTEMA]': str(spec_content.get("user_profiles", "")),
            '[LINK DO PROT√ìTIPO DE TELA]': str(spec_content.get("prototype_link", "N/A")),
            'Nome do Sistema': str(spec_content.get("system_name", "")),
            'NOME DO DOCUMENTO': str(spec_content.get("document_name", "")),
            'Alta': str(spec_content.get("importance", "")),
            'C√ìDIGO PROJETO': str(spec_content.get("project_code", "")),
            '00/00/0000': datetime.now().strftime("%d/%m/%Y"),
        }

        # Substitui placeholders no corpo do documento
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                     # Usar replace diretamente no texto do par√°grafo √© mais simples, 
                     # mas pode ter implica√ß√µes de estilo. Para este caso, deve funcionar.
                    p.text = p.text.replace(key, value)

        # Substitui placeholders nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, value)

        # Adiciona a se√ß√£o de funcionalidades dinamicamente
        functionalities_list = spec_content.get('functionalities', [])
        if functionalities_list:
            anchor_paragraph = None
            for p in doc.paragraphs:
                if '[TEXTO BREVE COM DESCRI√á√ÉO DA FUNCIONALIDADE]' in p.text:
                    anchor_paragraph = p
                    p.text = "" # Limpa o texto do placeholder
                    break
            
            if anchor_paragraph:
                # Verifica se o estilo de lista existe no documento
                has_list_bullet_style = False
                styles = doc.styles
                for s in styles:
                    if s.name == 'List Bullet':
                        has_list_bullet_style = True
                        break
                
                for i, func in enumerate(functionalities_list):
                    # Adiciona o t√≠tulo da funcionalidade
                    anchor_paragraph.insert_paragraph_before(f"6.{i+1} {func.get('title', 'Funcionalidade sem t√≠tulo')}", style='Heading 2')
                    
                    # Adiciona os detalhes
                    anchor_paragraph.insert_paragraph_before(f"Descri√ß√£o: {func.get('description', 'N/A')}")
                    anchor_paragraph.insert_paragraph_before(f"Acionador: {func.get('trigger', 'N/A')}")
                    anchor_paragraph.insert_paragraph_before(f"Integra√ß√µes: {func.get('integrations', 'N/A')}")
                    anchor_paragraph.insert_paragraph_before(f"V√≠nculo com Telas: {func.get('screen_links', 'N/A')}")
                    anchor_paragraph.insert_paragraph_before(f"Campos: {func.get('fields', 'N/A')}")
                    
                    # Adiciona os Requisitos Funcionais
                    rf_list = func.get('functional_requirements', [])
                    if rf_list:
                        anchor_paragraph.insert_paragraph_before("Requisitos Funcionais:")
                        for req in rf_list:
                            # L√ìGICA DE CORRE√á√ÉO:
                            # Se o estilo 'List Bullet' existir, usa. Sen√£o, adiciona um h√≠fen manualmente.
                            if has_list_bullet_style:
                                anchor_paragraph.insert_paragraph_before(req, style='List Bullet')
                            else:
                                anchor_paragraph.insert_paragraph_before(f"- {req}")
                    
                    # Adiciona um par√°grafo em branco para espa√ßamento
                    anchor_paragraph.insert_paragraph_before("") 

        return doc

    except Exception as e:
        print(f"Erro ao gerar documento Word: {e}")
        doc = Document()
        doc.add_paragraph(f"Ocorreu um erro ao gerar o documento Word: {e}")
        return doc